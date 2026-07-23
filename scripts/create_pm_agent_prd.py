from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("deliverables")
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "PM-Agent-MVP-PRD-v1.0.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "16324F"
INK = "1E293B"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F8FC"
WHITE = "FFFFFF"
GREEN = "276749"
AMBER = "8A5A00"
RED = "9B1C1C"
BORDER = "D8E0E8"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Title": (25, NAVY, 0, 4),
        "Subtitle": (13, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("PM AGENT  /  PRODUCT REQUIREMENTS DOCUMENT")
    set_run(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("PM Agent MVP PRD  |  Page ")
    set_run(run, size=9, color=MUTED)
    add_field(p, "PAGE")


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PM Agent MVP")
    set_run(r, size=28, bold=True, color=NAVY)

    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("From customer signal to engineering-ready work")
    set_run(r, size=14, color=MUTED)

    metadata = [
        ("Document status", "Draft for team review"),
        ("Version", "1.0"),
        ("Date", "17 July 2026"),
        ("Prepared for", "PM Agent Product and Engineering Team"),
        ("Product stage", "Deployable MVP / productionization required"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run(r, bold=True, color=NAVY)
        r = p.add_run(value)
        set_run(r, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_callout(
        doc,
        "Purpose of this PRD",
        "Align the team on the MVP user problem, product scope, functional requirements, production gaps, release criteria, and measurable outcomes. This document describes the current repository honestly: the experience is demo-ready and deployable, while persistence, real authentication, AI grounding, and several external integrations still require production implementation.",
        PALE_BLUE,
        BLUE,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Core product promise")
    set_run(r, size=12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "PM Agent converts scattered support signals into explainable, reviewable, engineering-ready tickets without removing human control."
    )
    set_run(r, size=15, bold=True, color=BLUE)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold_label=None, italic=False, color=INK):
    p = doc.add_paragraph()
    if bold_label and text.startswith(bold_label):
        r = p.add_run(bold_label)
        set_run(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_label):])
        set_run(r, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r)


def add_callout(doc, title, body, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=accent, size=8)
    repeat_table_header(table.rows[0])
    keep_row_together(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, bold=True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    repeat_table_header(header)
    keep_row_together(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run(r, size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        row = table.add_row()
        keep_row_together(row)
        set_table_geometry(table, widths_dxa)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run(r, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc):
    doc.add_page_break()


def build_doc():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_page(doc)

    add_page_break(doc)
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "PM Agent is a product operations workspace for Product Managers and Customer Support teams. It ingests customer requests, enriches them with product and code context, drafts structured specifications, and routes approved work toward engineering delivery. The product is designed to reduce manual ticket rewriting, shorten triage time, improve consistency, and preserve human accountability at every consequential decision.",
    )
    add_callout(
        doc,
        "MVP release definition",
        "A team can experience the full product journey in a deployed environment: sign in or onboard, connect representative sources, ask product questions, review AI-drafted tickets in Pipeline, apply ticket decisions, inspect insights, configure rules, and exercise available API routes. Production release requires the hardening work listed in Section 12.",
    )

    add_heading(doc, "1.1 Problem statement", 2)
    add_bullets(
        doc,
        [
            "Customer issues arrive across disconnected channels and are rewritten manually before engineering can act.",
            "Support agents often interrupt PMs to decide whether a report is expected behavior, a bug, or a feature request.",
            "Ticket quality varies: context, reproduction detail, scope, code references, and acceptance criteria are frequently incomplete.",
            "Existing queues show volume but do not consistently explain priority, confidence, or the next best action.",
            "Automation without an audit trail creates trust and accountability concerns.",
        ],
    )

    add_heading(doc, "1.2 Product vision", 2)
    add_para(
        doc,
        "Create one coherent workspace where teams can move from signal to decision to delivery, with grounded context, explicit guardrails, and a visible human approval step.",
    )

    add_heading(doc, "1.3 Product principles", 2)
    add_table(
        doc,
        ["Principle", "Product implication"],
        [
            ("Human control", "No consequential escalation or engineering handoff occurs without an explicit user action."),
            ("Explainability", "Drafts expose classification, confidence, rationale, evidence, and related code references."),
            ("Action over reporting", "Every primary screen should answer what the user should do next."),
            ("Traceability", "Decisions, edits, overrides, and automation outcomes must be attributable and reviewable."),
            ("Focused AI", "Ask PM is scoped to connected product context and must decline unsupported general conversation."),
        ],
        [2300, 7060],
    )

    add_page_break(doc)
    add_heading(doc, "2. Users and Jobs to Be Done", 1)
    add_heading(doc, "2.1 Primary users", 2)
    add_table(
        doc,
        ["User", "Primary job", "Key need"],
        [
            ("Product Manager", "Turn customer evidence into prioritized, reviewable engineering work.", "Fast decisions with context, confidence, and control."),
            ("Customer Support Agent", "Resolve product questions or escalate genuine issues without interrupting PMs.", "Grounded answers and a safe, confirmed ticket path."),
            ("Engineering / Dev Agent", "Receive a clear implementation-ready specification.", "Consistent scope, acceptance criteria, code references, and source context."),
            ("Workspace Admin", "Configure sources, repositories, rules, team access, and operational settings.", "Reliable setup, RBAC, auditability, and integration health."),
        ],
        [1900, 3940, 3520],
        font_size=9.2,
    )

    add_heading(doc, "2.2 Core jobs", 2)
    add_bullets(
        doc,
        [
            "When a customer report arrives, classify and scope it consistently so the right work enters the queue.",
            "When a support agent has a product question, answer from approved code and documentation or propose a confirmed escalation.",
            "When a PM reviews a draft, show enough evidence to accept, edit, reject, ignore, or investigate without switching tools.",
            "When work is approved, package it for the selected project tool or development agent with complete context.",
            "When performance changes, surface actionable insights and the underlying tickets rather than charts alone.",
        ],
    )

    add_heading(doc, "2.3 Roles and decision rights", 2)
    add_table(
        doc,
        ["Capability", "Owner/Admin", "User"],
        [
            ("Review and edit ticket drafts", "Allowed", "Allowed"),
            ("Accept, reject, or ignore tickets", "Allowed", "Allowed"),
            ("Configure integrations and automation", "Allowed", "Restricted by policy"),
            ("Invite or remove members", "Allowed", "Not allowed"),
            ("View audit history", "Allowed", "Read access as configured"),
        ],
        [5100, 2130, 2130],
    )

    add_page_break(doc)
    add_heading(doc, "3. Scope and Current Implementation", 1)
    add_para(
        doc,
        "The current repository is a Next.js 16 and React 19 MVP designed for Vercel. It includes production-style application routes and API handlers, but several core services still use mock data, browser persistence, or process-local server memory. The distinction below is mandatory for planning and stakeholder communication.",
    )
    add_table(
        doc,
        ["Area", "Current MVP state", "Production target"],
        [
            ("Authentication", "Sign-in, sign-up, reset, verification, and redirect UI; workspace state simulates session behavior.", "Managed identity provider, secure sessions, email verification, password reset, and server-side authorization."),
            ("Onboarding", "Guided source, repository, output, product-doc, and indexing experience using representative catalogs and delays.", "Real OAuth/API connections, durable credentials, indexing jobs, retries, and progress events."),
            ("Ask PM", "Global and ticket-context chat surfaces with scripted product responses and escalation behavior.", "Grounded retrieval over authorized repositories/docs, citations, safety boundaries, and evaluation monitoring."),
            ("Pipeline", "Review workspace, filters, clusters, ticket detail, reasoning, editing, decisions, and delivery states.", "Durable collaborative queue with concurrency control, audit history, and real downstream delivery."),
            ("Ingestion APIs", "Freshdesk and generic webhook routes normalize and create draft tickets in process-local memory.", "Authenticated, idempotent ingestion backed by a database and asynchronous processing."),
            ("Delivery / dispatch", "Payload generation, outbound dev-agent webhook, simulated external ticket URLs, and logged notification payloads.", "Real Jira/Linear/etc. adapters, email provider, retries, delivery status, and reconciliation."),
            ("Analytics / settings", "Functional UI over seeded and locally persisted state.", "Event-backed metrics, tenant-scoped configuration, durable audit logs, and policy enforcement."),
        ],
        [1700, 3710, 3950],
        font_size=8.6,
    )

    add_heading(doc, "3.1 In scope for the MVP experience", 2)
    add_bullets(
        doc,
        [
            "Authentication and onboarding user journeys.",
            "Ask PM global chat and ticket-context chat entry points.",
            "Ticket ingestion, classification, scope, confidence, code references, and human review actions.",
            "Canonical Pipeline workspace with ticket and cluster deep links.",
            "Connections, repositories, product documents, automation presets/rules, team, theme, and settings surfaces.",
            "Insights, activity, alerts, delivery stages, audit concepts, and downstream payload generation.",
            "Vercel-compatible build and runtime API routes.",
        ],
    )
    add_heading(doc, "3.2 Explicitly out of scope for this release", 2)
    add_bullets(
        doc,
        [
            "Autonomous code modification or pull-request creation by Ask PM.",
            "Silent ticket creation without user confirmation.",
            "General-purpose AI conversation outside connected product context.",
            "Claims of production-grade durability, security, or multi-tenant isolation before Section 12 gates are complete.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "4. End-to-End Product Flow", 1)
    add_numbers(
        doc,
        [
            "A user signs in or creates a workspace and completes onboarding.",
            "The workspace connects a customer-ticket source, one or more repositories, an output tool, and optional product documents.",
            "Inbound requests arrive through a connector or webhook and are normalized into the common ticket model.",
            "PM Agent classifies the request, estimates scope, identifies related code context, and creates a draft specification with confidence and reasoning.",
            "The draft appears in Pipeline, where a PM or authorized user reviews the source conversation, evidence, draft, and acceptance criteria.",
            "The reviewer accepts for development, accepts as non-technical, edits, rejects, ignores, or opens Ask PM with ticket context.",
            "Accepted engineering work is packaged for a project tool or dev agent; delivery status progresses from accepted to shipped.",
            "Events feed Insights, audit history, operational alerts, and improvement metrics.",
        ],
    )
    add_callout(
        doc,
        "Guardrail",
        "AI prepares and recommends; a human confirms. The product must retain source context, the generated draft, the final approved version, and the actor responsible for each decision.",
        "FFF8E8",
        AMBER,
    )

    add_heading(doc, "4.1 Primary navigation", 2)
    add_table(
        doc,
        ["Destination", "Purpose"],
        [
            ("Ask PM", "Ask grounded product questions or investigate a selected ticket."),
            ("Pipeline", "Review, prioritize, edit, and route tickets and clusters."),
            ("Connections", "Manage sources, repositories, outputs, and product context."),
            ("Automation", "Configure presets and readable if/then rules."),
            ("Settings", "Manage landing preference, notifications, workspace, profile, and product behavior."),
            ("Insights / Team", "Inspect outcomes, governance, roles, and accountability where enabled."),
        ],
        [2200, 7160],
    )

    add_page_break(doc)
    add_heading(doc, "5. Functional Requirements", 1)
    add_heading(doc, "5.1 Authentication and onboarding", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ("FR-001", "Users must be able to sign up, sign in, sign out, reset a password, and verify an email through a real identity service.", "Must"),
            ("FR-002", "New workspaces must complete a guided setup for ticket sources, repositories, output tools, and optional product documentation.", "Must"),
            ("FR-003", "Connection tests and indexing must report progress, completion, recoverable failures, and retry actions.", "Must"),
            ("FR-004", "The workspace must preserve onboarding completion and redirect authenticated users to their configured landing page.", "Must"),
        ],
        [1050, 7000, 1310],
        font_size=9.1,
    )

    add_heading(doc, "5.2 Ticket ingestion and enrichment", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ("FR-005", "The system must accept authenticated Freshdesk and generic webhook payloads and normalize them into one ticket schema.", "Must"),
            ("FR-006", "Duplicate source events must be detected using tenant, provider, and external ticket identifiers.", "Must"),
            ("FR-007", "Each draft must include classification, scope, description, suggested approach, acceptance criteria, confidence, rationale, and available code references.", "Must"),
            ("FR-008", "Original message, customer, source, attachments, notes, and conversation history must remain linked to the generated draft.", "Must"),
            ("FR-009", "Low-confidence or insufficient-context tickets must be flagged for manual review rather than silently completed.", "Must"),
        ],
        [1050, 7000, 1310],
        font_size=9.1,
    )

    add_heading(doc, "5.3 Ask PM", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ("FR-010", "Ask PM must support a global conversation and a ticket-context conversation that carries the selected ticket automatically.", "Must"),
            ("FR-011", "Answers must be grounded only in authorized repositories, product documents, and ticket history, with source citations where available.", "Must"),
            ("FR-012", "Simple questions may be answered directly; material issues must produce a proposed ticket summary and wait for confirmation.", "Must"),
            ("FR-013", "Ask PM must not edit code, create pull requests, or operate as an unrestricted general assistant.", "Must"),
            ("FR-014", "A ticket created from chat must retain the originating conversation and be labeled as chat-originated.", "Should"),
        ],
        [1050, 7000, 1310],
        font_size=9.1,
    )

    add_page_break(doc)
    add_heading(doc, "5.4 Pipeline review and decisions", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ("FR-015", "Pipeline must show pending tickets and clusters with filters for classification, scope, priority, confidence, age, and search.", "Must"),
            ("FR-016", "Selecting a ticket must show source context, editable draft fields, code references, confidence, and a reasoning trace without leaving the workspace.", "Must"),
            ("FR-017", "Authorized users must be able to accept for development, accept as non-technical, edit, reject, ignore, and undo eligible decisions.", "Must"),
            ("FR-018", "Cluster actions must support safe bulk review while retaining per-ticket history.", "Should"),
            ("FR-019", "Deep links must open the intended ticket or cluster and preserve supported query filters.", "Must"),
            ("FR-020", "Every decision and material edit must create an audit event with actor, time, object, and before/after context.", "Must"),
        ],
        [1050, 7000, 1310],
        font_size=9.1,
    )

    add_heading(doc, "5.5 Delivery, connections, automation, and insights", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ("FR-021", "Approved engineering work must generate a deterministic payload for the configured output tool or dev agent.", "Must"),
            ("FR-022", "Outbound delivery must expose queued, delivered, failed, retrying, and reconciled states; simulated URLs are not acceptable in production.", "Must"),
            ("FR-023", "Connections must display provider health, credential status, last successful sync, indexing state, and actionable errors.", "Must"),
            ("FR-024", "Automation must offer conservative, balanced, and aggressive presets plus readable if/then rules with preview before activation.", "Should"),
            ("FR-025", "Insights must connect metrics and recommendations to the underlying tickets and support delivery, quality, and chat-origin performance views.", "Should"),
            ("FR-026", "Settings must support theme, default landing page, notification preferences, workspace preferences, and tenant-scoped policy defaults.", "Should"),
        ],
        [1050, 7000, 1310],
        font_size=9.1,
    )

    add_page_break(doc)
    add_heading(doc, "6. Data, API, and Integration Requirements", 1)
    add_heading(doc, "6.1 Core data entities", 2)
    add_table(
        doc,
        ["Entity", "Minimum durable data"],
        [
            ("Workspace", "Tenant identity, plan, settings, lifecycle state, and retention policy."),
            ("User / membership", "Identity, workspace membership, role, status, and last access."),
            ("Integration", "Provider, type, encrypted credential reference, health, sync cursor, and error state."),
            ("Ticket", "Source identity, customer, conversation, classification, scope, draft, final decision, and timestamps."),
            ("Evidence", "Code/document reference, citation metadata, indexing version, and relevance score."),
            ("Delivery", "Destination, payload version, external ID/URL, state, attempts, and reconciliation result."),
            ("Audit event", "Actor, action, target, tenant, timestamp, correlation ID, and before/after summary."),
        ],
        [2300, 7060],
    )

    add_heading(doc, "6.2 Existing API surface", 2)
    add_table(
        doc,
        ["Endpoint group", "Purpose", "Production note"],
        [
            ("/api/webhooks/*", "Ingest Freshdesk or generic source events.", "Require authentication, signature validation, replay protection, idempotency, and rate limits."),
            ("/api/draft-tickets/*", "List/create drafts and record review decisions.", "Require tenant scoping, authorization, validation, persistence, and optimistic concurrency."),
            ("/api/product-docs", "List and register product documents.", "Add object storage, malware scanning, parsing jobs, versioning, and deletion policy."),
            ("/api/deliver", "Build downstream issue payload.", "Replace simulated external IDs/URLs with provider adapters and reconciliation."),
            ("/api/dispatch", "Send an approved ticket to a configured development-agent webhook.", "Use allow-listed destinations, signed requests, retry queue, timeout policy, and secret rotation."),
            ("/api/notify-dev", "Prepare developer notification content.", "Connect a real email/notification provider and delivery status tracking."),
        ],
        [2100, 3100, 4160],
        font_size=8.8,
    )

    add_heading(doc, "6.3 Environment and secrets", 2)
    add_bullets(
        doc,
        [
            "FRESHDESK_WEBHOOK_SECRET must be configured in deployed environments; an unset secret must not leave a production webhook open.",
            "FRESHDESK_API_KEY status alone is not a complete integration; credentials, domain/account identity, and actual API calls must be implemented securely.",
            "Provider credentials and webhook destinations must never be returned to the browser or stored in localStorage.",
            "Secrets must use Vercel environment configuration or an approved secret manager, with separate values by environment.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "7. Non-Functional Requirements", 1)
    add_table(
        doc,
        ["Category", "Requirement / target"],
        [
            ("Security", "Tenant isolation, server-side authorization, encrypted credentials, signed webhooks, least privilege, secure headers, abuse controls, and dependency scanning."),
            ("Reliability", "No accepted request is lost; ingestion and delivery are idempotent; failed jobs are retryable; durable backups and recovery procedures are tested."),
            ("Performance", "Primary authenticated pages should become interactive within 3 seconds at p75 on a normal broadband connection; common API reads should respond within 500 ms at p95 excluding external providers."),
            ("Observability", "Structured logs, request/correlation IDs, job metrics, integration health, error tracking, and alert thresholds must exist before production launch."),
            ("Accessibility", "Keyboard-operable workflows, visible focus, semantic controls, labeled forms, readable contrast, and WCAG 2.1 AA as the launch target."),
            ("Privacy", "Document data handling, retention, deletion, sub-processors, and code-indexing boundaries; avoid storing source content beyond approved policy."),
            ("Compatibility", "Support current stable desktop Chrome, Edge, Firefox, and Safari; responsive behavior for review and chat on small screens."),
            ("Maintainability", "Typed contracts, validation at boundaries, automated tests for critical flows, migration discipline, and environment-parity checks."),
        ],
        [2100, 7260],
        font_size=9.0,
    )

    add_heading(doc, "7.1 Audit and accountability", 2)
    add_bullets(
        doc,
        [
            "Store who performed each review action and whether it was manual or automated.",
            "Preserve the generated draft and the final approved content.",
            "Record rule changes, integration changes, role changes, dispatch attempts, and delivery outcomes.",
            "Make audit records append-only for normal users and exportable for administrators.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "8. Success Metrics", 1)
    add_para(
        doc,
        "The following are proposed targets for a controlled pilot. Baselines must be measured before interpreting improvement.",
    )
    add_table(
        doc,
        ["Outcome", "Metric", "Pilot target"],
        [
            ("Faster triage", "Median time from ingestion to first human decision", "50% lower than baseline"),
            ("Useful drafts", "Accepted or edit-and-accepted drafts / reviewed drafts", ">= 70%"),
            ("Lower review effort", "Median active review time per ticket", "<= 2 minutes"),
            ("Grounded answers", "Ask PM answers rated useful without escalation", ">= 75%"),
            ("Safe escalation", "Chat-created tickets with recorded user confirmation", "100%"),
            ("Reliable ingestion", "Valid source events processed successfully", ">= 99%"),
            ("Delivery quality", "Approved tickets delivered without manual payload repair", ">= 95%"),
            ("Trust", "Low-confidence drafts routed to manual review", "100%"),
        ],
        [2300, 4760, 2300],
    )

    add_heading(doc, "8.1 Instrumentation events", 2)
    add_bullets(
        doc,
        [
            "ticket_ingested, ticket_draft_created, ticket_review_opened, ticket_edited, ticket_decided",
            "chat_question_submitted, chat_answer_completed, chat_escalation_proposed, chat_ticket_confirmed",
            "delivery_started, delivery_succeeded, delivery_failed, delivery_retried",
            "integration_connected, integration_tested, integration_failed, indexing_started, indexing_completed",
            "automation_previewed, automation_enabled, automation_triggered, automation_overridden",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "9. MVP Acceptance Criteria", 1)
    add_heading(doc, "9.1 Demo acceptance", 2)
    add_bullets(
        doc,
        [
            "A clean installation completes the production build and starts successfully in a Vercel-compatible environment.",
            "A new user can traverse authentication, onboarding, and the configured landing route without broken navigation.",
            "A valid test webhook creates or identifies a draft ticket and the item is visible in the review workflow.",
            "A reviewer can open a deep-linked ticket, inspect evidence, edit the draft, make each supported decision, and observe the resulting state.",
            "Ask PM supports both global and ticket context and never performs an unconfirmed escalation.",
            "Connections, Automation, Settings, Insights, and Team surfaces render with coherent empty, loading, success, and error behavior.",
            "The primary light and dark themes remain readable and consistent across the core routes.",
        ],
    )

    add_heading(doc, "9.2 Quality acceptance", 2)
    add_bullets(
        doc,
        [
            "No build-blocking TypeScript errors, high-severity dependency vulnerabilities, or leaked secrets.",
            "Critical paths have automated tests for authorization, ingestion idempotency, review decisions, and delivery failure handling.",
            "No P0/P1 defects remain open; known P2 defects have documented owners and release decisions.",
            "Accessibility review covers keyboard navigation, forms, dialogs, contrast, focus, and screen-reader labels.",
            "Analytics events and error monitoring are verified in the target environment.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "10. Risks and Mitigations", 1)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ("Prototype state presented as production-ready", "Stakeholder trust, data loss, or security exposure.", "Use the release gates in Section 12 and label demo environments clearly."),
            ("In-memory or browser-only state", "Tickets, audits, and settings can disappear or diverge.", "Move tenant data to a transactional database with migrations and backups."),
            ("Ungrounded AI response", "Incorrect product guidance or poor engineering work.", "Use retrieval citations, confidence thresholds, evaluations, and human confirmation."),
            ("Weak webhook security", "Unauthorized ticket creation or replay attacks.", "Require signatures/secrets, timestamps, nonce/replay controls, rate limits, and idempotency."),
            ("External provider failure", "Stalled ingestion or delivery.", "Queue work, retry safely, show status, alert owners, and reconcile provider state."),
            ("Over-automation", "Incorrect tickets are accepted or dispatched.", "Conservative defaults, previews, bounded rules, approval thresholds, undo, and audit logs."),
            ("Sensitive code/document exposure", "Privacy and intellectual-property risk.", "Tenant isolation, access-scoped indexing, encryption, retention controls, and redaction."),
        ],
        [2600, 2800, 3960],
        font_size=8.8,
    )

    add_page_break(doc)
    add_heading(doc, "11. Delivery Plan", 1)
    add_table(
        doc,
        ["Phase", "Outcome", "Exit criteria"],
        [
            ("Phase 0 - Demo stabilization", "Deployable team demo with coherent core routes and representative data.", "Build green, smoke tests pass, no broken primary navigation, demo script approved."),
            ("Phase 1 - Durable foundation", "Real auth, tenant database, authorization, audit, secrets, and job infrastructure.", "Data survives deploys; cross-tenant tests pass; webhook and session security reviewed."),
            ("Phase 2 - Real integrations", "Freshdesk ingestion, repository indexing, product-doc processing, and Jira/Linear delivery.", "Sandbox end-to-end flows pass with retries, idempotency, and reconciliation."),
            ("Phase 3 - Grounded intelligence", "Evaluated classification, scoped retrieval, cited Ask PM answers, and confidence policy.", "Offline evaluation thresholds and pilot quality metrics are met."),
            ("Phase 4 - Controlled pilot", "Limited customer teams operate real workflows with support and monitoring.", "Pilot acceptance, security review, runbooks, backups, and go/no-go decision complete."),
        ],
        [2200, 3560, 3600],
        font_size=8.9,
    )

    add_heading(doc, "11.1 Recommended implementation order", 2)
    add_numbers(
        doc,
        [
            "Identity, workspace tenancy, RBAC, and durable database schema.",
            "Secure/idempotent ingestion plus background processing and audit events.",
            "Repository and product-document indexing with tenant-scoped retrieval.",
            "Grounded classification, drafting, Ask PM answers, confidence policy, and evaluations.",
            "Real output adapters, notifications, delivery reconciliation, and operational dashboards.",
            "Pilot hardening: tests, accessibility, security, privacy, performance, runbooks, and support readiness.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "12. Production Release Gates", 1)
    add_callout(
        doc,
        "Decision rule",
        "The deployed MVP may be shared as a product demonstration. It must not be described as production-ready until every Must gate below is verified in the production environment.",
        "FFF1F1",
        RED,
    )
    add_table(
        doc,
        ["Gate", "Required evidence", "Priority"],
        [
            ("Persistent data", "Database migrations, tenant-scoped schema, backups, restore test, and removal of process-local state for business records.", "Must"),
            ("Real authentication", "Identity provider, secure cookies/sessions, verification/reset flow, RBAC, and server-side access tests.", "Must"),
            ("Webhook security", "Configured secrets/signatures, timestamp validation, replay protection, idempotency, rate limits, and monitoring.", "Must"),
            ("AI grounding", "Authorized index, citations, evaluation set, confidence policy, refusal behavior, and prompt/data isolation tests.", "Must"),
            ("Real delivery", "At least one supported source and one supported output operate end to end without simulated URLs or logged-only notifications.", "Must"),
            ("Operational readiness", "Error tracking, dashboards, alerting, runbooks, owner rotation, incident process, and dependency/security scan.", "Must"),
            ("Quality", "Critical automated tests, accessibility review, performance baseline, cross-browser smoke tests, and no open P0/P1 issues.", "Must"),
            ("Privacy/legal", "Data map, retention/deletion behavior, customer disclosures, provider agreements, and code/document handling approval.", "Must"),
        ],
        [2200, 5820, 1340],
        font_size=8.9,
    )

    add_page_break(doc)
    add_heading(doc, "13. Open Decisions", 1)
    add_table(
        doc,
        ["Decision", "Why it matters", "Recommended owner"],
        [
            ("Initial pilot persona: PM-led or Support-led?", "Changes onboarding, default landing, metrics, and training emphasis.", "Product"),
            ("First production source and output providers", "Determines integration scope and pilot feasibility.", "Product + Engineering"),
            ("Repository indexing boundary and retention", "Defines privacy, cost, grounding quality, and customer approval needs.", "Security + Engineering"),
            ("Automation permission model", "Controls who can enable rules and which decisions may run automatically.", "Product + Security"),
            ("Confidence thresholds by classification", "Affects review volume, quality, and risk.", "Product + Data/AI"),
            ("Multi-workspace membership behavior", "Impacts tenancy, navigation, notifications, and billing.", "Product + Platform"),
        ],
        [3000, 4120, 2240],
        font_size=9.0,
    )

    add_heading(doc, "13.1 Definition of done for team approval", 2)
    add_bullets(
        doc,
        [
            "Product confirms the users, core promise, MVP scope, and success metrics.",
            "Engineering confirms the current-state assessment, production gaps, architecture direction, and sequencing.",
            "Design confirms the end-to-end journey, required states, accessibility target, and primary navigation.",
            "Security/privacy confirms the minimum controls for webhooks, credentials, code/doc indexing, retention, and audit.",
            "The team assigns owners and target dates to production release gates and open decisions.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Appendix A - Glossary", 1)
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ("Ask PM", "The grounded conversational surface for product questions and ticket investigation."),
            ("Draft ticket", "An AI-prepared specification awaiting a human decision."),
            ("Pipeline", "The canonical workspace for reviewing and tracking ticket work."),
            ("Classification", "The product category assigned to a request, such as bug, feature request, question, or churn signal."),
            ("Scope", "A relative implementation estimate: S, M, or L."),
            ("Reasoning trace", "The evidence and signals displayed to explain classification, confidence, and scope."),
            ("Code reference", "A linked repository file/function/line/snippet associated with a ticket."),
            ("Non-technical acceptance", "Resolution of a valid ticket without engineering delivery."),
            ("Dispatch", "Sending an approved specification to a development agent webhook."),
            ("Delivery", "Publishing approved work to the configured project or issue-tracking tool."),
        ],
        [2400, 6960],
    )

    add_heading(doc, "Appendix B - Source Basis", 1)
    add_para(
        doc,
        "This PRD was prepared from the current PM Agent repository, its implemented routes and domain model, the PM Agent 2.0 redesign proposal, the PM Agent 2.1 addendum, and the verified July 2026 deployment build. Where repository behavior is simulated or prototype-only, the document labels the production requirement separately.",
    )

    core = doc.core_properties
    core.title = "PM Agent MVP Product Requirements Document"
    core.subject = "Team-shareable PRD for the PM Agent MVP"
    core.author = "PM Agent Product Team"
    core.keywords = "PM Agent, PRD, product requirements, MVP"
    core.comments = "Prepared for internal product and engineering alignment."

    doc.save(DOCX_PATH)
    print(DOCX_PATH.resolve())


if __name__ == "__main__":
    build_doc()
