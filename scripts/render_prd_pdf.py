from html import escape
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    LongTable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfgen import canvas as pdf_canvas


INPUT = Path("deliverables/PM-Agent-MVP-PRD-v1.0.docx")
OUTPUT = Path("deliverables/PM-Agent-MVP-PRD-v1.0.pdf")

BLUE = colors.HexColor("#2E74B5")
DARK_BLUE = colors.HexColor("#1F4D78")
NAVY = colors.HexColor("#16324F")
INK = colors.HexColor("#1E293B")
MUTED = colors.HexColor("#64748B")
LIGHT_BLUE = colors.HexColor("#E8EEF5")
LIGHT_GRAY = colors.HexColor("#F2F4F7")
BORDER = colors.HexColor("#D8E0E8")
PALE_BLUE = colors.HexColor("#F4F8FC")
PAGE_WIDTH, PAGE_HEIGHT = letter
CONTENT_WIDTH = 6.5 * inch


def styles():
    base = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "PRDBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12.2,
            textColor=INK,
            spaceAfter=6,
            allowWidows=0,
            allowOrphans=0,
        ),
        "title": ParagraphStyle(
            "PRDTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=27,
            leading=31,
            textColor=NAVY,
            alignment=TA_LEFT,
            spaceBefore=8,
            spaceAfter=5,
        ),
        "subtitle": ParagraphStyle(
            "PRDSubtitle",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=13,
            leading=17,
            textColor=MUTED,
            spaceAfter=16,
        ),
        "h1": ParagraphStyle(
            "PRDH1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=BLUE,
            spaceBefore=13,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "PRDH2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=14,
            textColor=BLUE,
            spaceBefore=9,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "PRDH3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=DARK_BLUE,
            spaceBefore=7,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "bullet": ParagraphStyle(
            "PRDBullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.3,
            leading=11.8,
            textColor=INK,
            leftIndent=18,
            firstLineIndent=-10,
            bulletIndent=2,
            spaceAfter=5,
        ),
        "number": ParagraphStyle(
            "PRDNumber",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.3,
            leading=11.8,
            textColor=INK,
            leftIndent=22,
            firstLineIndent=-18,
            spaceAfter=5,
        ),
        "cell": ParagraphStyle(
            "PRDCell",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.8,
            leading=9.7,
            textColor=INK,
            spaceAfter=0,
        ),
        "cell_header": ParagraphStyle(
            "PRDCellHeader",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=7.9,
            leading=9.7,
            textColor=NAVY,
            spaceAfter=0,
        ),
        "callout_title": ParagraphStyle(
            "PRDCalloutTitle",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=12,
            textColor=BLUE,
            spaceAfter=3,
        ),
        "callout_body": ParagraphStyle(
            "PRDCalloutBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.1,
            leading=11.7,
            textColor=INK,
            spaceAfter=0,
        ),
    }


def paragraph_has_page_break(paragraph):
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def run_markup(paragraph):
    chunks = []
    for run in paragraph.runs:
        text = escape(run.text).replace("\n", "<br/>")
        if not text:
            continue
        if run.italic:
            text = f"<i>{text}</i>"
        if run.bold:
            text = f"<b>{text}</b>"
        chunks.append(text)
    return "".join(chunks) or escape(paragraph.text)


def table_widths(docx_table):
    grid_cols = docx_table._tbl.tblGrid.gridCol_lst
    widths = []
    for col in grid_cols:
        raw = col.get(qn("w:w"))
        widths.append(int(raw) if raw else 1)
    if not widths:
        widths = [1] * len(docx_table.columns)
    total = sum(widths)
    return [CONTENT_WIDTH * width / total for width in widths]


def make_callout(docx_table, s):
    cell = docx_table.cell(0, 0)
    paragraphs = [p for p in cell.paragraphs if p.text.strip()]
    title = paragraphs[0].text if paragraphs else "Note"
    body = " ".join(p.text for p in paragraphs[1:])
    content = [Paragraph(escape(title), s["callout_title"])]
    if body:
        content.append(Paragraph(escape(body), s["callout_body"]))
    box = Table([[content]], colWidths=[CONTENT_WIDTH], hAlign="LEFT")
    box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), PALE_BLUE),
                ("BOX", (0, 0), (-1, -1), 0.8, BLUE),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return [box, Spacer(1, 7)]


def make_table(docx_table, s):
    if len(docx_table.columns) == 1:
        return make_callout(docx_table, s)

    data = []
    for row_index, row in enumerate(docx_table.rows):
        row_data = []
        for cell in row.cells:
            text = "<br/>".join(escape(p.text) for p in cell.paragraphs if p.text.strip())
            row_data.append(Paragraph(text or " ", s["cell_header"] if row_index == 0 else s["cell"]))
        data.append(row_data)

    table = LongTable(
        data,
        colWidths=table_widths(docx_table),
        repeatRows=1,
        splitByRow=1,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT_BLUE),
                ("GRID", (0, 0), (-1, -1), 0.45, BORDER),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAFBFC")]),
            ]
        )
    )
    return [table, Spacer(1, 7)]


class HeaderFooterCanvas(pdf_canvas.Canvas):
    def _draw_header_footer(self):
        self.saveState()
        self.setFont("Helvetica-Bold", 7.5)
        self.setFillColor(MUTED)
        self.drawString(inch, PAGE_HEIGHT - 0.52 * inch, "PM AGENT  /  PRODUCT REQUIREMENTS DOCUMENT")
        self.setStrokeColor(BORDER)
        self.setLineWidth(0.45)
        self.line(inch, PAGE_HEIGHT - 0.61 * inch, PAGE_WIDTH - inch, PAGE_HEIGHT - 0.61 * inch)
        self.setFont("Helvetica", 8)
        self.setFillColor(MUTED)
        self.drawRightString(PAGE_WIDTH - inch, 0.55 * inch, f"PM Agent MVP PRD  |  Page {self._pageNumber}")
        self.restoreState()

    def showPage(self):
        self._draw_header_footer()
        super().showPage()


def build_pdf():
    s = styles()
    source = Document(INPUT)
    story = []
    list_number = 0

    for block in source.iter_inner_content():
        if isinstance(block, DocxTable):
            list_number = 0
            story.extend(make_table(block, s))
            continue

        if not isinstance(block, DocxParagraph):
            continue
        if paragraph_has_page_break(block):
            list_number = 0
            story.append(PageBreak())
            continue
        if not block.text.strip():
            story.append(Spacer(1, 3))
            continue

        style_name = block.style.name if block.style is not None else "Normal"
        markup = run_markup(block)
        if style_name == "Title":
            list_number = 0
            story.append(Paragraph(markup, s["title"]))
        elif style_name == "Subtitle":
            list_number = 0
            story.append(Paragraph(markup, s["subtitle"]))
        elif style_name == "Heading 1":
            list_number = 0
            story.append(Paragraph(markup, s["h1"]))
        elif style_name == "Heading 2":
            list_number = 0
            story.append(Paragraph(markup, s["h2"]))
        elif style_name == "Heading 3":
            list_number = 0
            story.append(Paragraph(markup, s["h3"]))
        elif style_name == "List Bullet":
            list_number = 0
            story.append(Paragraph(markup, s["bullet"], bulletText="\u2022"))
        elif style_name == "List Number":
            list_number += 1
            story.append(Paragraph(f"{list_number}.&nbsp;&nbsp;{markup}", s["number"]))
        else:
            list_number = 0
            story.append(Paragraph(markup, s["body"]))

    pdf = SimpleDocTemplate(
        str(OUTPUT),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.78 * inch,
        bottomMargin=0.78 * inch,
        title="PM Agent MVP Product Requirements Document",
        author="PM Agent Product Team",
        subject="Team-shareable PRD for the PM Agent MVP",
    )
    pdf.build(story, canvasmaker=HeaderFooterCanvas)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_pdf()
