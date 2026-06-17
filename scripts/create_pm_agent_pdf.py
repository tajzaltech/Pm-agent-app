from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("deliverables")
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "PM-Agent-How-It-Works.docx"


TITLE = "PM Agent: How the Tool Works"
SUBTITLE = "A simple English guide for understanding the full product flow"


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string("2E4BFF" if level == 1 else "243B53")
        run.font.size = Pt(16 if level == 1 else 13)
        run.bold = True
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_para(doc, text, bold_start=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    if bold_start and text.startswith(bold_start):
        first = para.add_run(bold_start)
        set_run(first, bold=True)
        rest = para.add_run(text[len(bold_start):])
        set_run(rest)
    else:
        run = para.add_run(text)
        set_run(run)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_run(run)
    return para


def add_number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_run(run)
    return para


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    cell.text = ""
    shading = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6FF")
    shading.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, bold=True, color="2E4BFF")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(TITLE)
    set_run(r, size=24, bold=True, color="2E4BFF")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run(SUBTITLE)
    set_run(r, size=12, color="555555")

    add_note(
        doc,
        "Short summary",
        "PM Agent turns customer support requests into clear engineering tickets. It connects to ticket sources, reads code context from repositories, drafts implementation-ready tickets, and lets the team approve or reject them before sending them to a project management tool.",
    )

    add_heading(doc, "1. What PM Agent Is", 1)
    add_para(
        doc,
        "PM Agent is a product management assistant for software teams. It helps teams convert customer requests, bugs, and support messages into organized development work.",
    )
    add_para(
        doc,
        "The tool does not replace a product manager or engineer. It prepares a first draft, adds useful context, and gives the team a review screen so a human can make the final decision.",
    )
    add_bullet(doc, "Support teams use it to avoid manually rewriting tickets.")
    add_bullet(doc, "Product teams use it to understand customer demand.")
    add_bullet(doc, "Engineering teams use it to see related code areas and implementation notes.")

    add_heading(doc, "2. First User Flow", 1)
    add_para(doc, "The normal user journey is simple:")
    add_number(doc, "The user opens the app and first sees the sign in page.")
    add_number(doc, "If the user already has an account, they sign in and go directly to the main app.")
    add_number(doc, "If the user is new, they click sign up.")
    add_number(doc, "After sign up, the user goes to onboarding.")
    add_number(doc, "Onboarding connects tools and prepares the workspace.")
    add_number(doc, "After setup, the user lands in the Queue to review AI-created tickets.")

    add_note(
        doc,
        "Current prototype behavior",
        "In the current demo build, sign in goes to the Queue, and sign up goes directly to onboarding. In a real production app, sign up can also include email verification before onboarding.",
    )

    add_heading(doc, "3. Authentication Screens", 1)
    add_para(doc, "The authentication area contains the pages a normal SaaS product needs:")
    add_bullet(doc, "Sign in: for existing users.")
    add_bullet(doc, "Sign up: for new users creating a workspace.")
    add_bullet(doc, "Forgot password: lets a user request a reset link.")
    add_bullet(doc, "Reset password: lets the user choose a new password.")
    add_bullet(doc, "Verify email: can be used when the product requires email confirmation.")
    add_para(
        doc,
        "The purpose of these screens is to make the product feel like a real application before the user reaches onboarding and the main dashboard.",
    )

    add_heading(doc, "4. Onboarding Flow", 1)
    add_para(
        doc,
        "Onboarding teaches PM Agent where information comes from, where the code lives, and where accepted tickets should be sent.",
    )

    add_heading(doc, "Step 1: Connect Ticket Source", 2)
    add_para(
        doc,
        "The ticket source is where customer requests come from. Examples include Freshdesk, Zendesk, Jira Service Management, Salesforce, Google Sheets, or a custom webhook.",
    )
    add_bullet(doc, "Freshdesk or Zendesk can bring support tickets into PM Agent.")
    add_bullet(doc, "Salesforce can bring customer or CRM-related requests.")
    add_bullet(doc, "Google Sheets can be used for simple imported request lists.")
    add_bullet(doc, "A webhook lets any external system send ticket data.")

    add_heading(doc, "Step 2: Connect Code Repository", 2)
    add_para(
        doc,
        "The repository connection tells PM Agent where the software code lives. The app supports tools such as GitHub and Bitbucket.",
    )
    add_para(
        doc,
        "PM Agent uses the repository to understand structure, routes, functions, models, and files that may relate to a customer request.",
    )
    add_note(
        doc,
        "Important privacy idea",
        "The intended product behavior is to create a semantic map of the codebase. The user interface states that source code itself is not stored; instead, PM Agent keeps useful code context for matching tickets to implementation areas.",
    )

    add_heading(doc, "Step 3: Connect Output Tool", 2)
    add_para(
        doc,
        "The output tool is where approved development tickets should go. Examples include Linear, Jira, Monday.com, ClickUp, and GitHub Issues.",
    )
    add_para(
        doc,
        "This step matters because PM Agent should not only analyze requests. It should also push accepted work into the team's normal delivery system.",
    )

    add_heading(doc, "Step 4: Upload Product Docs", 2)
    add_para(
        doc,
        "Product docs give PM Agent more context. The user can upload documentation, requirements, guides, or product notes.",
    )
    add_bullet(doc, "Docs help the agent understand product language.")
    add_bullet(doc, "Docs help reduce vague or incomplete ticket drafts.")
    add_bullet(doc, "Docs make acceptance criteria more accurate.")

    add_heading(doc, "5. Indexing", 1)
    add_para(
        doc,
        "After onboarding, PM Agent indexes the workspace. Indexing means the app reads and organizes the connected information so it can be searched and matched later.",
    )
    add_para(doc, "During indexing, the app can show progress such as:")
    add_bullet(doc, "Reading repository structure.")
    add_bullet(doc, "Parsing functions and routes.")
    add_bullet(doc, "Analyzing models and schemas.")
    add_bullet(doc, "Building a semantic map.")
    add_bullet(doc, "Completing setup.")

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "6. Main App: Queue", 1)
    add_para(
        doc,
        "The Queue is the main working page. It shows draft engineering tickets created from customer requests. The team reviews these drafts before pushing them to the selected output tool.",
    )
    add_para(doc, "Each queue row usually shows:")
    add_bullet(doc, "Ticket title.")
    add_bullet(doc, "Customer or source information.")
    add_bullet(doc, "Classification such as bug, feature request, question, or churn signal.")
    add_bullet(doc, "Scope estimate such as small, medium, or large.")
    add_bullet(doc, "Related code references.")
    add_bullet(doc, "Actions like accept, reject, or view.")

    add_heading(doc, "7. Ticket Review Page", 1)
    add_para(
        doc,
        "The ticket review page gives a detailed side-by-side view. One side shows the original customer request. The other side shows the AI-generated development ticket.",
    )
    add_para(doc, "The customer side can include:")
    add_bullet(doc, "Original subject and message.")
    add_bullet(doc, "Customer name, email, and plan.")
    add_bullet(doc, "Conversation history.")
    add_bullet(doc, "Internal notes.")
    add_bullet(doc, "Attachments.")
    add_para(doc, "The AI draft side can include:")
    add_bullet(doc, "Draft title.")
    add_bullet(doc, "Description.")
    add_bullet(doc, "Classification and scope.")
    add_bullet(doc, "Affected code references.")
    add_bullet(doc, "Suggested technical approach.")
    add_bullet(doc, "Acceptance criteria.")

    add_heading(doc, "8. Human Review Actions", 1)
    add_para(
        doc,
        "PM Agent is designed around human approval. The AI creates a draft, but the user decides what happens next.",
    )
    add_bullet(doc, "Accept: sends or marks the ticket as approved for the output tool.")
    add_bullet(doc, "Reject: removes the draft from the active review flow.")
    add_bullet(doc, "Edit and accept: lets the user improve the title, description, scope, or acceptance criteria before approval.")
    add_bullet(doc, "Undo: allows the user to quickly reverse a recent action.")

    add_heading(doc, "9. Other App Areas", 1)
    add_para(doc, "The sidebar includes other product areas that support the review workflow:")
    add_bullet(doc, "Dashboard: shows high-level performance and recent activity.")
    add_bullet(doc, "Clusters: groups related tickets or repeated customer problems.")
    add_bullet(doc, "Analytics: helps understand ticket trends and product demand.")
    add_bullet(doc, "Connections: shows connected tools and integration status.")
    add_bullet(doc, "Settings: manages preferences, team settings, and rules.")
    add_bullet(doc, "Logout: returns the user to the sign in page.")

    add_heading(doc, "10. Example Scenario", 1)
    add_para(
        doc,
        "Imagine a customer writes to support: 'Users are being charged twice during checkout.' PM Agent can process that request and create a draft engineering ticket.",
    )
    add_number(doc, "The request enters from Freshdesk or Zendesk.")
    add_number(doc, "PM Agent checks the connected repository for payment-related files.")
    add_number(doc, "It finds likely code areas such as webhook handlers or payment services.")
    add_number(doc, "It creates a ticket title, description, suggested approach, and acceptance criteria.")
    add_number(doc, "The product or engineering team reviews the draft in the Queue.")
    add_number(doc, "If the draft looks correct, the team accepts it and sends it to Linear, Jira, or another output tool.")

    add_heading(doc, "11. What Makes the Tool Useful", 1)
    add_bullet(doc, "It saves time by turning raw support messages into structured dev tickets.")
    add_bullet(doc, "It gives engineers better starting context.")
    add_bullet(doc, "It helps product teams identify repeated pain points.")
    add_bullet(doc, "It keeps the final decision with a human reviewer.")
    add_bullet(doc, "It connects customer evidence to delivery work.")

    add_heading(doc, "12. Simple Mental Model", 1)
    add_para(doc, "You can understand PM Agent with this simple flow:")
    add_note(
        doc,
        "Input -> Context -> AI Draft -> Human Review -> Delivery Tool",
        "Customer tickets come in, PM Agent adds product and code context, AI prepares a draft, a human reviews it, and accepted work goes to the team's project management tool.",
    )

    add_heading(doc, "13. Notes for Real Production Use", 1)
    add_para(
        doc,
        "The current app is a polished product prototype. For a production version, the following backend features would normally be needed:",
    )
    add_bullet(doc, "Real authentication and secure sessions.")
    add_bullet(doc, "Real OAuth connections to ticket sources, repositories, and output tools.")
    add_bullet(doc, "Secure storage for integration tokens.")
    add_bullet(doc, "Background jobs for indexing and AI analysis.")
    add_bullet(doc, "Database persistence for users, workspaces, tickets, and actions.")
    add_bullet(doc, "Audit logs and permission controls.")

    add_heading(doc, "14. Final Summary", 1)
    add_para(
        doc,
        "PM Agent is a workflow tool for software teams. It starts with sign in or sign up, guides the user through onboarding, connects support tools and repositories, creates AI-assisted ticket drafts, and lets the team review everything before sending work into delivery tools.",
    )
    add_para(
        doc,
        "The main value is not only AI writing. The value is combining customer context, code context, product documentation, and human approval in one clear workflow.",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
